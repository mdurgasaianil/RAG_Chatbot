import os
import uuid
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import tempfile

from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# Load environment variables
load_dotenv()
os.environ["GEMINI_API_KEY"] = os.getenv("GEMINI_API_KEY")

# Initialize FastAPI app
app = FastAPI(
    title="DocMind RAG Chatbot",
    description="Upload any document and chat with it",
    version="1.0.0"
)

# CORS — allows Gowthami's React app to call this API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # in production replace with her Vercel URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Gemini LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0.3
)

print("App initialized ✓")

def load_document(file_path: str) -> str:
    """
    Loads text from PDF, DOCX, or TXT file
    """
    ext = file_path.split(".")[-1].lower()
    
    if ext == "pdf":
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    
    elif ext == "docx":
        doc = DocxDocument(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    elif ext == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    else:
        raise ValueError(f"Unsupported file type: {ext}")

# Session storage — stores each user's data in memory
sessions = {}

class SessionData:
    def __init__(self):
        self.vector_store: Chroma = None
        self.chat_history: list = []
        self.file_name: str = ""

def get_session(session_id: str) -> SessionData:
    """
    Gets existing session or raises error if not found
    """
    if session_id not in sessions:
        raise HTTPException(
            status_code=404,
            detail=f"Session {session_id} not found. Please upload a document first."
        )
    return sessions[session_id]

def create_session() -> str:
    """
    Creates a new session and returns its ID
    """
    session_id = str(uuid.uuid4())
    sessions[session_id] = SessionData()
    return session_id

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """
    Accepts a document, processes it, stores in ChromaDB
    Returns a session_id for future chat requests
    """
    # Step 1: Validate file type
    allowed_types = ["pdf", "docx", "txt"]
    ext = file.filename.split(".")[-1].lower()
    if ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Allowed: {allowed_types}"
        )

    # Step 2: Save uploaded file temporarily to disk
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Step 3: Load document text (Cell 3 logic)
        raw_text = load_document(tmp_path)
        if not raw_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Document appears to be empty"
            )

        # Step 4: Chunk the text (Cell 4 logic)
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", ".", " "]
        )
        chunks = splitter.split_text(raw_text)
        documents = [Document(page_content=chunk) for chunk in chunks]

        # Step 5: Create embeddings + vector store (Cell 5 logic)
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-001"
        )
        vector_store = Chroma.from_documents(
            documents=documents,
            embedding=embeddings,
            collection_name=f"session_{str(uuid.uuid4())[:8]}"
        )

        # Step 6: Create session and store vector store
        session_id = create_session()
        sessions[session_id].vector_store = vector_store
        sessions[session_id].file_name = file.filename

        return {
            "session_id": session_id,
            "file_name": file.filename,
            "total_chunks": len(documents),
            "message": "Document uploaded and processed successfully"
        }

    finally:
        # Step 7: Always delete the temp file
        os.unlink(tmp_path)

# Request body model
class ChatRequest(BaseModel):
    session_id: str
    question: str

async def generate_stream(question: str, vector_store: Chroma, session_id: str):
    """
    Generator function that streams answer token by token
    """
    # Step 1: Build retriever
    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 8}
    )

    # Step 2: Build prompt
    prompt = ChatPromptTemplate.from_template("""
    You are a helpful assistant. Answer the question based ONLY on the context provided below.
    If the answer is not in the context, say "I don't have enough information to answer this."
    
    Context:
    {context}
    
    Question:
    {question}
    
    Answer:
    """)

    # Step 3: Format docs helper
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    # Step 4: Build RAG chain
    rag_chain = (
        {
            "context": retriever | format_docs,
            "question": RunnablePassthrough()
        }
        | prompt
        | llm
        | StrOutputParser()
    )

    # Step 5: Stream tokens one by one in SSE format
    full_answer = ""
    async for chunk in rag_chain.astream(question):
        full_answer += chunk
        yield f"data: {chunk}\n\n"  # SSE format

    # Step 6: Save to chat history after streaming completes
    sessions[session_id].chat_history.append({
        "question": question,
        "answer": full_answer
    })

    # Step 7: Send done signal to frontend
    yield "data: [DONE]\n\n"


@app.post("/chat")
async def chat(request: ChatRequest):
    """
    Receives question, streams answer back via SSE
    """
    # Validate session exists
    session = get_session(request.session_id)

    # Validate question
    if not request.question.strip():
        raise HTTPException(
            status_code=400,
            detail="Question cannot be empty"
        )

    return StreamingResponse(
        generate_stream(
            question=request.question,
            vector_store=session.vector_store,
            session_id=request.session_id
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"  # prevents nginx from buffering the stream
        }
    )

@app.get("/history/{session_id}")
async def get_history(session_id: str):
    """
    Returns all past Q&A for a session
    """
    session = get_session(session_id)
    
    return {
        "session_id": session_id,
        "file_name": session.file_name,
        "total_messages": len(session.chat_history),
        "history": session.chat_history
    }


@app.delete("/session/{session_id}")
async def delete_session(session_id: str):
    """
    Clears a session and its vector store from memory
    """
    session = get_session(session_id)
    
    # Delete ChromaDB collection
    session.vector_store.delete_collection()
    
    # Remove session from memory
    del sessions[session_id]
    
    return {
        "message": f"Session {session_id} deleted successfully"
    }


@app.delete("/sessions/all")
async def delete_all_sessions():
    """
    Clears ALL sessions — useful for dev/testing
    """
    count = len(sessions)
    for session_id in list(sessions.keys()):
        sessions[session_id].vector_store.delete_collection()
    sessions.clear()
    
    return {
        "message": f"Cleared {count} sessions successfully"
    }


@app.get("/health")
async def health_check():
    """
    Health check endpoint — confirms API is running
    """
    return {
        "status": "running",
        "active_sessions": len(sessions),
        "model": "gemini-2.5-flash"
    }